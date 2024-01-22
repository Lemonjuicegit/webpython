import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT,WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

def setCelltext(table_,row_,cell_,text_,font_size_=Pt(10.5)):
    table_.rows[row_].cells[cell_].paragraphs[0].text = str(text_)
    table_.rows[row_].cells[cell_].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table_.rows[row_].cells[cell_].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table_.rows[row_].cells[cell_].paragraphs[0].runs[0].font.size = font_size_

def addAreaStatistics(table_,data_={'DKH':'','ZDDM':'','MJ':''}):
    row_len_ = len(table_.rows)
    cell_ = table_.cell
    table_.add_row()
    table_.add_row()
    table_.add_row()
    table_.add_row()
    table_.add_row()
    table_.add_row()
    table_.add_row()
    setCelltext(table_,row_len_,0,data_['DKH'])
    setCelltext(table_,row_len_,1,f"{data_['ZDDM'][:10]}\n{data_['ZDDM'][10:]}")
    setCelltext(table_,row_len_,2,round(data_['MJ']))
    setCelltext(table_,row_len_,3,'农用地')
    setCelltext(table_,row_len_+1,3,'其中')
    setCelltext(table_,row_len_+1,4,'耕地')
    setCelltext(table_,row_len_+2,4,'林地')
    setCelltext(table_,row_len_+3,4,'草地')
    setCelltext(table_,row_len_+4,4,'其他')
    setCelltext(table_,row_len_+5,3,'建设用地')
    setCelltext(table_,row_len_+6,3,'未利用地')
    setCelltext(table_,row_len_,5,data_['LYD'] if 'LYD' in data_ else '0')
    setCelltext(table_,row_len_+1,5,data_['GD'] if 'GD' in data_ else '0')
    setCelltext(table_,row_len_+2,5,data_['LD'] if 'LD' in data_ else '0')
    setCelltext(table_,row_len_+3,5,data_['CD'] if 'CD' in data_ else '0')
    setCelltext(table_,row_len_+4,5,data_['QT'] if 'QT' in data_ else '0')
    setCelltext(table_,row_len_+5,5,data_['JSYD'] if 'JSYD' in data_ else '0')
    setCelltext(table_,row_len_+6,5,data_['WLYD'] if 'WLYD' in data_ else '0')
    
    cell_(row_len_,0).merge(cell_(row_len_+6,0))
    cell_(row_len_,1).merge(cell_(row_len_+6,1))
    cell_(row_len_,2).merge(cell_(row_len_+6,2))
    cell_(row_len_+1,3).merge(cell_(row_len_+4,3))
    cell_(row_len_,3).merge(cell_(row_len_,4))
    cell_(row_len_+5,3).merge(cell_(row_len_+5,4))
    cell_(row_len_+6,3).merge(cell_(row_len_+6,4))
    
def addZDStatistics(table_,data_={'DKH':'','ZDDM':'','MJ':''}):
    row_len_ = len(table_.rows)
    cell_ = table_.cell
    table_.add_row()
    cell_(row_len_,1).merge(cell_(row_len_,2))
    cell_(row_len_,3).merge(cell_(row_len_,5))
    setCelltext(table_,row_len_,0,data_['DKH'])
    setCelltext(table_,row_len_,1,data_['ZDDM'])
    setCelltext(table_,row_len_,3,round(data_['MJ'],4))

def collectivelyZDClassifiedArea(data,savePath):
  test_doc = Document()
  test_doc.styles['Normal'].font.name = u'方正仿宋_GBK'
  test_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),u'方正仿宋_GBK')
  test_doc.styles['Normal'].font.size = Pt(10.5)
  test_doc.add_table(6,6,style="Table Grid")
  table = test_doc.tables[0]
  table.alignment = WD_TABLE_ALIGNMENT.CENTER
  cell = table.cell
  colu = table.columns
  table.rows[0].cells[0].paragraphs[0].text = '集体土地所有权宗地分类面积调查表'
  table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(15)
  table.rows[0].cells[0].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

  table.rows[1].cells[5].paragraphs[0].text = '单位：m2'
  table.rows[1].cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

  setCelltext(table,2,0,'权利人')
  setCelltext(table,2,2,'重庆市沙坪坝区凤凰镇威灵寺村催家坪社')
  setCelltext(table,3,0,'宗地代码')
  setCelltext(table,3,2,data['ZDDM_TXST'])
  setCelltext(table,4,0,'不动产单元号')
  setCelltext(table,4,2,data['BDCDYH_TXST'])
  setCelltext(table,5,0,'地块号')
  setCelltext(table,5,1,'宗地代码')
  setCelltext(table,5,2,'面积（m2）')
  setCelltext(table,5,3,'分类面积（m2）')
  cell(5,3).merge(cell(5,5))

  cell(3,0).merge(cell(3,1))
  cell(3,2).merge(cell(3,5))

  table.rows[4].cells[0].paragraphs[0].text = '不动产单元号'
  table.rows[4].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
  table.rows[4].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
  cell(4,0).merge(cell(4,1))
  cell(4,2).merge(cell(4,5))

  cell(0,0).merge(cell(1,5))
  cell(2,0).merge(cell(2,1))
  cell(2,2).merge(cell(2,5))

  colu[0].width = Cm(3.03)
  colu[1].width = Cm(2.61)
  colu[2].width = Cm(2.12)
  colu[3].width = Cm(1.99)
  colu[4].width = Cm(1.91)
  colu[5].width = Cm(3.56)
  table.rows[0].height = Cm(1.48)
  table.rows[2].height = Cm(1.26)
  for data_row in data['area_statistics']:
    addAreaStatistics(table,data['area_statistics'][data_row])

  if data['fd_data']:
    table.add_row()
    setCelltext(table,len(table.rows)-1,0,'其他社飞入本社面积',Pt(14))
    cell(len(table.rows)-1,0).merge(cell(len(table.rows)-1,5))
    addZDStatistics(table,{'DKH':'地块号','ZDDM':'宗地代码','MJ':'面积（m2）'})
    for data_row in data['fd_data']:
        addZDStatistics(table,data['fd_data'][data_row])
    table.add_row()
    setCelltext(table,len(table.rows)-1,0,'合计面积')
    setCelltext(table,len(table.rows)-1,3,round(data['fd_total_area'],4))
    cell(len(table.rows)-1,0).merge(cell(len(table.rows)-1,2))
    cell(len(table.rows)-1,3).merge(cell(len(table.rows)-1,5))
  test_doc.save(savePath)

def get_data(df,fddata_=''):
  datadf = {}
  for index, row in df.iterrows(): 
    if row['地类'] in ['耕地','林地','草地','其他农用地','建设用地','未利用地',]:
      if row['权利人'] not in datadf:
        datadf[row['权利人']] = {
          'ZDDM_TXST':row['宗地代码'],
          'BDCDYH_TXST':row['不动产单元号'],
          'fd_data':{},
          'fd_total_area':0,
          'area_statistics':{
            row['宗地代码']:{
              'DKH':index,
              'MJ':row['面积'],
              'ZDDM':row['宗地代码'],
              'LYD':row['面积'] if row['地类'] in ['耕地','林地','草地','其他农用地'] else 0,
              'GD':row['面积'] if row['地类'] == '耕地' else 0,
              'LD':row['面积'] if row['地类'] == '林地' else 0,
              'CD':row['面积'] if row['地类'] == '草地' else 0,
              'QT':row['面积'] if row['地类'] == '其他农用地' else 0,
              'JSYD':row['面积'] if row['地类'] == '建设用地' else 0,
              'WLYD':row['面积'] if row['地类'] == '未利用地' else 0
            }
          }
        }
      else:
        if row['宗地代码'] not in datadf[row['权利人']]['area_statistics']:   
          datadf[row['权利人']]['ZDDM_TXST'] = f"{datadf[row['权利人']]['ZDDM_TXST']}、{row['宗地代码']}"
          datadf[row['权利人']]['BDCDYH_TXST'] = f"{datadf[row['权利人']]['BDCDYH_TXST']}、{row['不动产单元号']}"
          datadf[row['权利人']]['area_statistics'][row['宗地代码']] = {
            'DKH':index,
            'ZDDM':row['宗地代码'],
            'LYD':row['面积'] if row['地类'] in ['耕地','林地','草地','其他农用地'] else 0,
            'MJ':row['面积'],
            'GD':row['面积'] if row['地类'] == '耕地' else 0,
            'LD':row['面积'] if row['地类'] == '林地' else 0,
            'CD':row['面积'] if row['地类'] == '草地' else 0,
            'QT':row['面积'] if row['地类'] == '其他农用地' else 0,
            'JSYD':row['面积'] if row['地类'] == '建设用地' else 0,
            'WLYD':row['面积'] if row['地类'] == '未利用地' else 0
          }
        else:
          temp = datadf[row['权利人']]['area_statistics'][row['宗地代码']]
          datadf[row['权利人']]['area_statistics'][row['宗地代码']] = {
                      'DKH':index,
                      'ZDDM':row['宗地代码'],
                      'MJ':row['面积']+temp['MJ'],
                      'LYD':row['面积']+temp['MJ'] if row['地类'] in ['耕地','林地','草地','其他农用地'] else temp['LYD'],
                      'GD':row['面积'] if row['地类'] == '耕地' else temp['GD'],
                      'LD':row['面积'] if row['地类'] == '林地' else temp['LD'],
                      'CD':row['面积'] if row['地类'] == '草地' else temp['CD'],
                      'QT':row['面积'] if row['地类'] == '其他农用地' else temp['QT'],
                      'JSYD':row['面积'] if row['地类'] == '建设用地' else temp['JSYD'],
                      'WLYD':row['面积'] if row['地类'] == '未利用地' else temp['WLYD'],
                  }
  if fddata_:
    df_fd = pd.read_excel(fddata_)
    for index, row in df_fd.iterrows():
      datadf[row['权利人']]['fd_data'][row['宗地代码']] = {'DKH':index,'ZDDM':row['宗地代码'],'MJ':row['面积']}
      if row['权利人'] in datadf:
        datadf[row['权利人']]['fd_total_area'] = row['面积'] + datadf[row['权利人']]['fd_total_area'] 
      else:
        datadf[row['权利人']]['fd_total_area'] = row['面积']
  return datadf

def Area_table_all(df,df_fda,savepath):
  # @df_fda:飞地数据
  for key,value in get_data(df,df_fda).items():
    collectivelyZDClassifiedArea(value,Path(savepath) / f"{key}分类面积调查表.docx")
    yield f"{key}分类面积调查表"

if __name__ == '__main__':
  pass

