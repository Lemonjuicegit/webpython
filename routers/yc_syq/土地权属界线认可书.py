import re
from docx import Document
from docxtpl import DocxTemplate
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
from . import config
from . import Djmod
log = Djmod.Djlog()

template_path = Path(config.config.template_path)

# docx文档赋值
def setCelltext(table_,
                row_,
                cell_,
                text_,
                paragraph_=0,
                run_=0,
                font_name_='',
                font_size_=Pt(10.5)):
    """_summary_
    word单元格居中赋值
  Args:
      table_ (table): python-docx模块table对象
      row_ (number): 行号
      cell_ (number): 列号
      text_ (str/number): 需要赋值的内容
      paragraph_ (int, optional): 段落号 Defaults to 0.
      run_ (int, optional): 样式号 Defaults to 0.
      font_name_ (str, optional): 字体名称. Defaults to ''.
      font_size_ (_type_, optional): 字体大小 Defaults to Pt(10.5).
  """
    table_.rows[row_].cells[cell_].paragraphs[paragraph_].add_run(str(text_))
    table_.rows[row_].cells[
        cell_].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table_.rows[row_].cells[cell_].paragraphs[
        paragraph_].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table_.rows[row_].cells[cell_].paragraphs[paragraph_].runs[
        run_].font.size = font_size_
    if font_name_:
        r = table_.rows[row_].cells[cell_].paragraphs[paragraph_].runs[run_]
        r.font.name = font_name_
        r._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_)

def stamp_paralist(sift_jzx):
    result = {}
    for _,row in sift_jzx.iterrows():
        if row['LZQLRMC']:
            if row['LZQLRMC'] not in result:
                result[row['LZQLRMC']] = {
                    'jzdh_para':f"{row['QSDH']}至{row['ZZDH']}({row['ZDDM']})",
                    'is':True if row["LZQLRMC"][-4:] in ['村民小组','农民集体'] else False,
                    'LZQLRMC':row['LZQLRMC'],
                    'XLXZ':row['XLXZ'],
                    'XLCM':row['XLCM'],
                    'XLSM':row['XLSM'],
                }
            else:
                result[row['LZQLRMC']]['jzdh_para'] = f"{result[row['LZQLRMC']]['jzdh_para']}、{row['QSDH']}至{row['ZZDH']}({row['ZDDM']})"
    return result

def generate_jxrks(qlr,gdf_ZD,gdf_jzx):
    doc = DocxTemplate(template_path/'土地权属界线认可书.docx')
    sift_jzx = gdf_jzx[gdf_jzx.QLRMC == qlr]
    if not sift_jzx.size:
        log.err(f"界线认可书——{qlr}:没有界址线")
        return
    XLXZ = list(filter(None,sift_jzx.XLXZ.drop_duplicates().values))
    XLCM = list(filter(None,sift_jzx.XLCM.drop_duplicates().values))
    XLSM = list(filter(None,sift_jzx.XLSM.drop_duplicates().values))
    XLXZ = [v for v in XLXZ if v != '']
    XLCM = [v for v in XLCM if v != '']
    XLSM = [v for v in XLSM if v not in ['','桐梁区','大足区','璧山区','江津区','荣昌区']]
    doc.render({
        'BXZ':sift_jzx['BXZ'].values[0],
        'BCM':sift_jzx['BCM'].values[0],
        'BSM':sift_jzx['BSM'].values[0],
        'QXLSM':XLSM[0],
        'XLXZ':'、'.join(XLXZ) if XLXZ else '',
        'XLCM':'、'.join(XLCM) if XLCM else '',
        'XLSM':'、'.join(XLSM) if XLSM else '',
        'ZXLSM':'、'.join(XLSM[1:-1]) if XLSM else '',
        'JZLSM':XLSM[-1],
        'TFH':'、'.join(gdf_ZD[gdf_ZD.QLRMC == qlr]['TFH'].drop_duplicates().values),
        'stamp_paralist':[value for value in stamp_paralist(sift_jzx).values()]        
    })
    return doc
         
def get_jxzx(jzd_jzsm,jzx_jzsm):
    if not jzx_jzsm.size:
        return ""
    result = f"{jzx_jzsm.QLRMC.values[0]}{jzx_jzsm.ZDDM.values[0]}\n"
    for _, row in jzx_jzsm.iterrows():
        if row["LZQLRMC"]:
            JZXLB = jzd_jzsm[jzd_jzsm["JZD_NEW"] == row["QSDH"]]["JZXLB"].values[0]
            H_JZXLB = jzd_jzsm[jzd_jzsm["JZD_NEW"] == row["ZZDH"]]["JZXLB"].values[0]
            QSDH = row["QSDH"]
            LZQLRMC = row["LZQLRMC"]
            ZJDH = f"过界址点{row['ZJDH']}" if row["ZJDH"] else ""
            ZZDH = row["ZZDH"]
            result += f"{QSDH}沿着{LZQLRMC}{JZXLB}{ZJDH}到{ZZDH}止，后为{H_JZXLB}。\n"
    return result

def generate_jzsmjb(gdf_ZD,gdf_jzd,qlr):
    doc = Document(template_path / '界址说明标示表.docx') # type: ignore
    data_ = gdf_jzd.copy()
    doc.styles['Normal'].font.name = u'方正仿宋_GBK'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋_GBK')
    tfh = '、'.join(gdf_ZD[gdf_ZD.QLRMC == qlr].TFH.drop_duplicates().values)
    cell = doc.tables[0].cell
    setCelltext(doc.tables[0], 0, 4, qlr)
    setCelltext(doc.tables[0], 1, 4, tfh)
    countlist = [c for c in Djmod.groupby(data_,['ZDDM','INDEX'],'count').COUNT]
    data_['ZDDM_INDEX'] = data_.ZDDM.str.cat(data_.INDEX.astype(str))
    zddm_index = data_['ZDDM_INDEX'].drop_duplicates()
    count = len(countlist) * 2
    for c in countlist:
        if c > 6:
            count += 12
        else:
            count += c * 2
    if count < 30:
        count = 30
    for i in range(count):
        doc.tables[0].add_row()
        doc.tables[0].rows[i + 4].height = Cm(0.5)
    for row_index in range(0, count, 2):
        if row_index <= count - 4:
            cell(row_index + 5, 0).merge(cell(row_index + 6, 0))
        for colu_index in range(8):
            cell(row_index + 4, colu_index + 1).merge(
                cell(row_index + 5, colu_index + 1)
            )
    num = 0
    JZXLB_dict = {'沟渠':1,'道路':2,'田埂':3,'地埂':4,'山脊':5}
    JZXWZ_dict = {'内':6,'中':7,'外':8}
    for z_i in zddm_index:
        zddm = z_i[:19]
        jdz_boundary = data_[data_.ZDDM_INDEX == z_i].reset_index()
        if jdz_boundary.shape[0] == 0:
            continue
        if jdz_boundary.shape[0] <= 6:
            for _,row in jdz_boundary.iterrows():
                setCelltext(doc.tables[0], num + 4, 0, row['JZD_NEW'])
                if row['JZXLB'] in JZXLB_dict:
                    setCelltext(doc.tables[0], num + 5, JZXLB_dict[row['JZXLB']], '√')
                else:
                    log.err(f"{zddm}-{row['JZD_NEW']}:界址线类别填写不正确")    

                if row['JZXWZ'] in JZXWZ_dict:
                    setCelltext(doc.tables[0], num + 5, JZXWZ_dict[row['JZXWZ']], '√')
                else:
                    log.err(f"{zddm}-{row['JZD_NEW']}:界址线位置填写不正确") 
                num += 2 

            setCelltext(doc.tables[0], num + 4, 0, jdz_boundary.at[0,'JZD_NEW'])
            num += 2
        else:
            for i in range(6):
                if i < 3:
                    JZD_NEW =   jdz_boundary.at[i,'JZD_NEW']
                    JZXLB   =   jdz_boundary.at[i,'JZXLB']
                    JZXWZ   =   jdz_boundary.at[i,'JZXWZ']
                    setCelltext(doc.tables[0], num + 4, 0,JZD_NEW)
                    if JZXLB in JZXLB_dict:
                        setCelltext(doc.tables[0], num + 5, JZXLB_dict[JZXLB], '√')
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线类别填写不正确")
                        
                    if JZXWZ in JZXWZ_dict:
                        setCelltext(doc.tables[0], num + 5, JZXWZ_dict[JZXWZ], '√')
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线位置填写不正确")     

                    num += 2
                elif i == 4:
                    setCelltext(doc.tables[0], num + 4, 0, '. . .')
                    setCelltext(doc.tables[0], num + 4, 4, '√')
                    setCelltext(doc.tables[0], num + 4, 7, '√')
                    num += 2
                elif i > 4:
                    JZD_NEW =   jdz_boundary.iloc[-1].JZD_NEW
                    JZXLB   =   jdz_boundary.iloc[-1].JZXLB
                    JZXWZ   =   jdz_boundary.iloc[-1].JZXWZ
                    setCelltext( doc.tables[0], num + 4, 0,JZD_NEW)
                    
                    if JZXLB in JZXLB_dict:
                        setCelltext(doc.tables[0], num + 5, JZXLB_dict[JZXLB], '√')
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线类别填写不正确")
                        
                    if JZXWZ in JZXWZ_dict:
                        setCelltext(doc.tables[0], num + 5, JZXWZ_dict[JZXWZ], '√')
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线位置填写不正确")     
                    num += 2
            setCelltext(doc.tables[0], num + 4, 0, jdz_boundary.at[0,'JZD_NEW'])
            num += 2
    # for zddm in ZDDMLIST:
    #     dwsm = ''  # 点位说明
    #     jxzx = ''  # 界线走向
    #     jzd_jzsm = gdf_jzd[gdf_jzd['ZDDM'] == zddm]
    #     jzx_jzsm = gdf_jzx[gdf_jzx['ZDDM'] == zddm]
    #     dwsm += f"{qlr}{zddm}\n"
    #     dwsmdict = {}
    #     for _, row in jzd_jzsm.iterrows():
    #         if row['点位说明']:
    #             if row['点位说明'] not in dwsmdict:
    #                 dwsmdict[row['点位说明']] = row['JZD_NEW']
    #             else:
    #                 dwsmdict[row['点位说明']] = f"{dwsmdict[row['点位说明']]}、{row['JZD_NEW']}"
    #     for key,value in dwsmdict.items():
    #         dwsm += f"{value}位于{key}交界处。\n"
    #     print(zddm)
    #     jxzx += get_jxzx(jzd_jzsm,jzx_jzsm)
    # setCelltext(doc.tables[1], 0, 1, dwsm)
    # setCelltext(doc.tables[1], 1, 1, jxzx)
    return doc

def generate_jzsm(zddmlist, qlr, jzd_data, jzx_data):
    doc = DocxTemplate(template_path / "界址说明表认可书.docx")
    dwsm = ""  # 点位说明
    jxzx = ""  # 界线走向
    for zddm in zddmlist:
        log.info(f"{zddm}的界址信息")
        jzd_jzsm = jzd_data[jzd_data["ZDDM"] == zddm]
        jzx_jzsm = jzx_data[jzx_data["ZDDM"] == zddm]
        dwsmdict = {}
        if [v for v in jzd_jzsm.点位说明.values if v]:
            dwsm += f"{qlr}{zddm}\n"
        for _, row in jzd_jzsm.iterrows():
            if row["点位说明"]:
                if row["点位说明"] not in dwsmdict:
                    dwsmdict[row["点位说明"]] = row["JZD_NEW"]
                else:
                    dwsmdict[row["点位说明"]] = f"{dwsmdict[row['点位说明']]}、{row['JZD_NEW']}"
        for key, value in dwsmdict.items():
            dwsm += f"{value}位于{key}交界处。\n"
        jxzx += get_jxzx(jzd_jzsm, jzx_jzsm)
    doc.render({"DWSM": dwsm, "JXZX": jxzx})
    return doc

def generate_jxrks_all(gdf_ZD,gdf_jzd,jzx_df,savepath,control):
    qlrs = gdf_ZD['QLRMC'].drop_duplicates().values
    for qlr in qlrs:
        doclist = []
        ZDDM_list = gdf_ZD[gdf_ZD.QLRMC == qlr].ZDDM.drop_duplicates().values
        jzjb_data = gdf_jzd[gdf_jzd.ZDDM.isin(ZDDM_list)]
        if control['jxrks']:
            jxrks = generate_jxrks(qlr,gdf_ZD,jzx_df)
            if not jxrks:
                log.info(f"{qlr}未出线认可书")
                yield f"{qlr}未出线认可书"
                continue
            doclist.append(jxrks)
        if control['jzsmjb']:
            jzsmjb = generate_jzsmjb(gdf_ZD,jzjb_data,qlr)
            doclist.append(jzsmjb)
            jzsm = generate_jzsm(ZDDM_list, qlr, jzjb_data, jzx_df)
            doclist.append(jzsm)
        Djmod.compose_docx(doclist, Path(savepath)/ f"{qlr}界线认可书.docx") # type: ignore
        yield f"{qlr}界线认可书"

if __name__ == '__main__':
    pass
        