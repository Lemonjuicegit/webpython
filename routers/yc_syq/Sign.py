import os
from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

def docxtabel_indaex(docxtabel):
  row = docxtabel.rows
  s = 0
  for i in range(len(row)):
    cell = row[i].cells
    d = 0
    for r in cell:
      print(f"{s}-{d}:{r.text}")
      d += 1
    s += 1

def docxpar_indaex(docxpar):
  row = docxpar.paragraphs
  s = 0
  for i in range(len(row)):
    cell = row[i].runs
    d = 0
    for r in cell:
      print(f"{s}-{d}:{r.text}")
      d += 1
    s += 1

#   修改段落内容
def makevalue(doc_templet, paragraph, runs, value, string):
    para = doc_templet.paragraphs
    para[paragraph].runs[runs].text = para[paragraph].runs[runs].text.replace(value, string)

#   修改doc文档插入表格的内容
def maketablevalue(doc_templet, table, columu, row, paragraph, runs,value, string):
  # 带原有格式修改值
  doc_templet.tables[table].cell(columu, row).paragraphs[paragraph].runs[runs].text = \
    doc_templet.tables[table].cell(columu, row).paragraphs[paragraph].runs[runs].text.replace(value, string)

def maketableparagraphs(doc_templet, table, columu, row, paragraph,value, string):
    doc_templet.tables[table].cell(columu, row).paragraphs[paragraph].text = \
        doc_templet.tables[table].cell(columu, row).paragraphs[paragraph].text.replace(value, string)

#   向docx文档插入的表格添加内容
def addtablevalue(doc_templet, table, columu, row, paragraph, value):  # 向
    doc_templet.tables[table].cell(columu, row).paragraphs[paragraph].text = value

def excel_dict(xlsx_path,sheet_name):
  execl = load_workbook(xlsx_path, data_only=True)
  sheet = execl[sheet_name]
  lis = []
  dic = []
  for row in sheet.rows:
    temp = []
    for value in row:
      temp.append(value.value)
    lis.append(temp)
  index = lis[0]
  del lis[0]
  for row in lis:
    temp = {}
    for i in range(len(row)):
      temp[index[i]] = row[i]
    dic.append(temp)
  return dic

def real_estate_sq(date,save_path):
  # 不动产登记申请表
  # 填充内容位置：权利人姓名（2-3）、坐落（14-3）、不动产单元号（15-3）、面积（17-3）
  doc_sq = Document('./template/不动产登记申请表.docx')
  maketablevalue(doc_sq,0,2,3,0,0,'[qlr]',date['QSDWMC'])
  maketablevalue(doc_sq,0,14,3,0,0,'[zl]',date['镇村社'])
  maketablevalue(doc_sq,0,15,3,0,0,'[bdcdyh]',date['不动产单元号'])
  doc_sq.tables[0].cell(15,3).paragraphs[0].runs[0].font.size = Pt(9)
  maketablevalue(doc_sq,0,17,3,0,0,'[mj]',str(round(date['面积'],0)))
  save_path_ = os.path.join(save_path,date['村名'],'签字资料',date['QSDWMC'])
  if not os.path.exists(save_path_):
    os.makedirs(save_path_)
  doc_sq.save(os.path.join(save_path_,'不动产登记申请表.docx'))


def real_estate_frwt(date,save_path):
  # 法人授权委托书
  # 填充内容位置：单位全称(0-1),单位地址(1-1),受理单位(14-2),
  doc_temp = Document('./template/法人授权委托书.docx')
  makevalue(doc_temp,14,2,'[sldw]',date['QSDWMC'])
  maketablevalue(doc_temp,0,0,1,0,0,'[dwmc]',date['QSDWMC'])
  maketablevalue(doc_temp,0,1,1,0,0,'[dwdz]',date['QSDWMC'])
  maketablevalue(doc_temp,1,1,4,0,0,'[dwmc]',date['QSDWMC'])
  maketablevalue(doc_temp,1,1,5,0,0,'[dwdz]',date['QSDWMC'])
  doc_temp.tables[1].cell(1,4).paragraphs[0].runs[0].font.size = Pt(9)
  doc_temp.tables[1].cell(1,5).paragraphs[0].runs[0].font.size = Pt(9)
  save_path_ = os.path.join(save_path,date['村名'],'签字资料',date['QSDWMC'])
  if not os.path.exists(save_path_):
    os.makedirs(save_path_)
  doc_temp.save(os.path.join(save_path_,'法人授权委托书.docx'))
  
def real_estate_zjrzm(date,save_path):
  # 指界人证明
  # 1-6,1-9,1-13,4-4,4-9,4-13,
  doc_temp = Document('./template/指界人证明.docx')
  makevalue(doc_temp,1,6,'[xzm]',date['乡镇名'])
  makevalue(doc_temp,1,9,'[cm]',date['村名'])
  makevalue(doc_temp,1,13,'[sm]',date['社名'])
  makevalue(doc_temp,4,4,'[xzm]',date['乡镇名'])
  makevalue(doc_temp,4,9,'[cm]',date['村名'])
  makevalue(doc_temp,4,13,'[sm]',date['社名'])

  save_path_ = os.path.join(save_path,date['村名'],'签字资料',date['QSDWMC'])
  if not os.path.exists(save_path_):
    os.makedirs(save_path_)
  doc_temp.save(os.path.join(save_path_,'指界人证明.docx'))

def real_estate_zjtzs(date,save_path):
  # 指界通知书
  # 3-3,5-2,9-0,10-26,11-2
  doc_temp = Document('./template/指界通知书.docx')
  makevalue(doc_temp,3,3,'[zl]',date['QSDWMC'])
  makevalue(doc_temp,5,2,'[jhdd]',date['镇村社'])
  makevalue(doc_temp,9,0,'[qlr]',date['QSDWMC'])
  makevalue(doc_temp,10,28,'[qlr]',date['QSDWMC'])
  makevalue(doc_temp,11,2,'[jhdd]',date['镇村社'])

  save_path_ = os.path.join(save_path,date['村名'],'签字资料',date['QSDWMC'])
  if not os.path.exists(save_path_):
    os.makedirs(save_path_)
  doc_temp.save(os.path.join(save_path_,'指界通知书.docx'))

def real_estate_zjwts(date,save_path):
  # 指界委托书
  # 2-16
  doc_temp = Document('./template/指界委托书.docx')
  makevalue(doc_temp,2,16,'[qlr]',date['QSDWMC'])
  makevalue(doc_temp,12,1,'[txdz]',date['txdz'])

  save_path_ = os.path.join(save_path,date['村名'],'签字资料',date['QSDWMC'])
  if not os.path.exists(save_path_):
    os.makedirs(save_path_)
  doc_temp.save(os.path.join(save_path_,'指界委托书.docx'))
def real_estate_frzms(date,save_path):
  # 法人证明书
  # 15-2
  doc_temp = Document('./template/法人证明书.docx')
  makevalue(doc_temp,15,2,'[qsdwmc]',date['QSDWMC'])
  save_path_ = os.path.join(save_path,date['村名'],'签字资料',date['QSDWMC'])
  if not os.path.exists(save_path_):
    os.makedirs(save_path_)
  doc_temp.save(os.path.join(save_path_,'法人证明书.docx'))

def real_estate_yssq(date,save_path):
  # 遗失申请
  # 3-1 3-5
  doc_temp = Document('./template/遗失申明.docx')
  makevalue(doc_temp,3,1,'[qlr]',date['QSDWMC'])
  makevalue(doc_temp,3,5,'[zl]',date['QSDWMC'])
  save_path_ = os.path.join(save_path,date['村名'],'签字资料',date['QSDWMC'])
  if not os.path.exists(save_path_):
    os.makedirs(save_path_)
  doc_temp.save(os.path.join(save_path_,'遗失申明.docx'))

def get_date(xlsx_teb):
  date =  excel_dict(xlsx_teb,'Sheet1')
  date_dic = {}
  for row in date:
    if row['QSDWMC'] in date_dic:
      temp = date_dic[row['QSDWMC']]['不动产单元号']
      date_dic[row['QSDWMC']]['不动产单元号'] = f"{temp}、{row['不动产单元号']}"
      date_dic[row['QSDWMC']]['面积'] = date_dic[row['QSDWMC']]['面积'] + row['宗地面积']
    else:
      date_dic[row['QSDWMC']] = {'社名':row['社名'],'乡镇名':row['乡镇名'],'村名':row['村名'],'QSDWMC':row['QSDWMC'],'镇村社':row['镇村社'],'不动产单元号':row['不动产单元号'],'面积':row['宗地面积'],'txdz':row['txdz']}
  return date_dic

def main(date_dic,save_path):
  n = 1
  for row in date_dic:
    yield f"正在生成{row}资料: {n}/{len(date_dic)}"
    real_estate_sq(date_dic[row],save_path)
    real_estate_frwt(date_dic[row],save_path)
    real_estate_zjrzm(date_dic[row],save_path)
    real_estate_zjtzs(date_dic[row],save_path)
    real_estate_zjwts(date_dic[row],save_path)
    real_estate_frzms(date_dic[row],save_path)
    real_estate_yssq(date_dic[row],save_path)
    n += 1

    