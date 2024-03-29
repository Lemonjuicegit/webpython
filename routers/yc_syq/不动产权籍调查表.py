import re
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from . import Djmod
from pathlib import Path
from . import config


template_path = Path(config.config.template_path)
log = Djmod.Djlog()


def jpg_pathlist(jpg_zdct):
    ct_path = {}
    for file in jpg_zdct.glob("*宗地草图.jpg"):
        ct_path[file.name[:-8]] = str(file)
    return ct_path


# docx文档赋值
def setCelltext(
    table_, row_, cell_, text_, paragraph_=0, run_=0, font_name_="", font_size_=Pt(10.5)
):
    """_summary_
      word单元格居中赋值
    Args:
        table_ (table): python-docx模块table对象
        row_ (number): 行号
        cell_ (number): 列号
        text_ (str/number): 需要赋值的内容
        paragraph_ (int, optional): 段落号 Defaults to 0.
        run_ (int, optional): 样式号 Defaults to 0.
        font_name_ (str, optional): 字体名称. Defaults to ''.
        font_size_ (_type_, optional): 字体大小 Defaults to Pt(10.5).
    """
    table_.rows[row_].cells[cell_].paragraphs[paragraph_].add_run(str(text_))
    table_.rows[row_].cells[
        cell_
    ].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table_.rows[row_].cells[cell_].paragraphs[
        paragraph_
    ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table_.rows[row_].cells[cell_].paragraphs[paragraph_].runs[
        run_
    ].font.size = font_size_
    if font_name_:
        r = table_.rows[row_].cells[cell_].paragraphs[paragraph_].runs[run_]
        r.font.name = font_name_
        r._element.rPr.rFonts.set(qn("w:eastAsia"), font_name_)


# 宗地数据格式化
def get_zd_data(zd_data, jzx_data, jzd_gdf):
    tidyup = {}
    n = 0
    # 宗地数据格式化
    for _, row in zd_data.iterrows():
        if row["QLRMC"] not in tidyup:
            tidyup[row["QLRMC"]] = {
                **row,
                "NUM": 1,
                "jzx_data": [],
                "ZDDM_list": {row["ZDDM"]},
            }
        else:
            tidyup[row["QLRMC"]]["ZDDM"] = f"{tidyup[row['QLRMC']]['ZDDM']}、{row['ZDDM']}"
            tidyup[row["QLRMC"]]["ZDMJ"] = tidyup[row["QLRMC"]]["ZDMJ"] + row["ZDMJ"]
            tidyup[row["QLRMC"]][
                "BDCDYH"
            ] = f"{tidyup[row['QLRMC']]['BDCDYH']}、{row['BDCDYH']}"
            if row["TFH"]:
                temp_tfh = set(row["TFH"].split("、"))
                tidyup_tfh = set(tidyup[row["QLRMC"]]["TFH"].split("、"))
                over = tidyup_tfh | temp_tfh
                tidyup[row["QLRMC"]]["TFH"] = "、".join(over)
            tidyup[row["QLRMC"]]["ZDDM_list"].add(row["ZDDM"])
            tidyup[row["QLRMC"]]["NUM"] = tidyup[row["QLRMC"]]["NUM"] + 1
    for _, row in jzx_data.iterrows():
        try:
            xlzl = jzd_gdf[jzd_gdf["ZDDM"] == row["ZDDM"]]
            tidyup[row["QLRMC"]]["jzx_data"].append(
                {
                    "QLRMC": row["QLRMC"],
                    "QSDH": row["QSDH"],
                    "ZJDH": row["ZJDH"],
                    "ZZDH": row["ZZDH"],
                    "LZQLRMC": row["LZQLRMC"] if row["LZQLRMC"] else "",
                    "qz1": f"详见{row['LZQLRMC']}界线认可书" if row["LZQLRMC"][-4:] in ['村民小组','农民集体'] else "",
                    "qz2": f"详见{row['QLRMC']}界线认可书" if row["QLRMC"] else "",
                    "dkm": xlzl.loc[xlzl.index[0], "SM"]
                    if xlzl.loc[xlzl.index[0], "SM"]
                    else "",
                }
            )
        except KeyError:
            log.err(f"界址线与宗地不匹配: {row['ZDDM']}")
    # tidyup = {'双石镇丁家岩村楠竹林村民小组':tidyup['双石镇丁家岩村楠竹林村民小组']}
    return tidyup


def generate_head(zddm):
    doc = Document(template_path / "不动产权籍调查表.docx")  # type: ignore
    doc.paragraphs[8].runs[2].text = (
        doc.paragraphs[8].runs[2].text.replace("@ZDDM@", zddm)
    )
    return doc


def generate_zd(data_):
    doc = DocxTemplate(template_path / "宗地基本信息表.docx")
    doc.render(data_)
    return doc


def generate_jzjb(jzddf):
    # data_:{int:[{}]}
    data_ = jzddf.copy()
    doc = Document(template_path / "界址标示表.docx")  # type: ignore
    doc.styles["Normal"].font.name = "方正仿宋_GBK"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "方正仿宋_GBK")
    cell = doc.tables[0].cell
    countlist = [c for c in Djmod.groupby(data_, ["ZDDM", "INDEX"], "count").COUNT]
    data_["ZDDM_INDEX"] = data_.ZDDM.str.cat(data_.INDEX.astype(str))
    zddm_index = data_["ZDDM_INDEX"].drop_duplicates()
    count = len(countlist) * 2
    for c in countlist:
        if c > 6:
            count += 12
        else:
            count += c * 2
    if count < 32:
        count = 32
    for i in range(count):
        doc.tables[0].add_row()
        doc.tables[0].rows[i + 3].height = Cm(0.5)
    for row_index in range(0, count, 2):
        if row_index <= count - 4:
            cell(row_index + 4, 0).merge(cell(row_index + 5, 0))
            cell(row_index + 4, 1).merge(cell(row_index + 5, 1))
        for colu_index in range(10):
            cell(row_index + 3, colu_index + 2).merge(
                cell(row_index + 4, colu_index + 2)
            )
    num = 0
    JZXLB_dict = {"沟渠": 3, "道路": 4, "田埂": 5, "地埂": 6, "山脊": 7}
    JZXWZ_dict = {"内": 8, "中": 9, "外": 10}

    for z_i in zddm_index:
        zddm = z_i[:19]
        jdz_boundary = data_[data_.ZDDM_INDEX == z_i].reset_index()
        if jdz_boundary.shape[0] == 0:
            continue
        if jdz_boundary.shape[0] <= 6:
            for _, row in jdz_boundary.iterrows():
                setCelltext(doc.tables[0], num + 3, 0, row["JZD_NEW"])
                if row["JZXLB"] in JZXLB_dict:
                    setCelltext(doc.tables[0], num + 4, JZXLB_dict[row["JZXLB"]], "√")
                else:
                    log.err(f"{zddm}-{row['JZD_NEW']}:界址线类别填写不正确")

                if row["JZXWZ"] in JZXWZ_dict:
                    setCelltext(doc.tables[0], num + 4, JZXWZ_dict[row["JZXWZ"]], "√")
                else:
                    log.err(f"{zddm}-{row['JZD_NEW']}:界址线位置填写不正确")

                if row["SM"]:
                    setCelltext(doc.tables[0], num + 4, 11, row["SM"])
                    doc.tables[0].cell(num + 4, 11).paragraphs[0].runs[
                        0
                    ].font.size = Pt(7.5)
                num += 2
            setCelltext(doc.tables[0], num + 3, 0, jdz_boundary.at[0, "JZD_NEW"])
            num += 2
        else:
            for i in range(6):
                if i < 3:
                    JZD_NEW = jdz_boundary.at[i, "JZD_NEW"]
                    JZXLB = jdz_boundary.at[i, "JZXLB"]
                    JZXWZ = jdz_boundary.at[i, "JZXWZ"]
                    SM = jdz_boundary.at[i, "SM"]
                    setCelltext(doc.tables[0], num + 3, 0, JZD_NEW)
                    if JZXLB in JZXLB_dict:
                        setCelltext(doc.tables[0], num + 4, JZXLB_dict[JZXLB], "√")
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线类别填写不正确")

                    if JZXWZ in JZXWZ_dict:
                        setCelltext(doc.tables[0], num + 4, JZXWZ_dict[JZXWZ], "√")
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线位置填写不正确")

                    if SM:
                        setCelltext(doc.tables[0], num + 4, 11, SM)
                        doc.tables[0].cell(num + 4, 11).paragraphs[0].runs[
                            0
                        ].font.size = Pt(7.5)
                    num += 2
                elif i == 4:
                    setCelltext(doc.tables[0], num + 3, 0, ". . .")
                    setCelltext(doc.tables[0], num + 3, 6, "√")
                    setCelltext(doc.tables[0], num + 3, 9, "√")
                    num += 2
                elif i > 4:
                    JZD_NEW = jdz_boundary.iloc[-1].JZD_NEW
                    JZXLB = jdz_boundary.iloc[-1].JZXLB
                    JZXWZ = jdz_boundary.iloc[-1].JZXWZ
                    SM = jdz_boundary.iloc[-1].SM
                    setCelltext(doc.tables[0], num + 3, 0, JZD_NEW)
                    if JZXLB in JZXLB_dict:
                        setCelltext(doc.tables[0], num + 4, JZXLB_dict[JZXLB], "√")
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线类别填写不正确")
                    if JZXWZ in JZXWZ_dict:
                        setCelltext(doc.tables[0], num + 4, JZXWZ_dict[JZXWZ], "√")
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线位置填写不正确")
                    if SM:
                        setCelltext(doc.tables[0], num + 4, 11, SM)
                        doc.tables[0].cell(num + 4, 11).paragraphs[0].runs[
                            0
                        ].font.size = Pt(7.5)
                    num += 2
            setCelltext(doc.tables[0], num + 3, 0, jdz_boundary.at[0, "JZD_NEW"])
            num += 2
    return doc


# 界址签章表组件
# def add_jzqz(table, data):
#     table.add_row()
#     setCelltext(table,len(table.rows) - 1,0, data['QSDH'],font_name_='Times New Roman')
#     if data['ZJDH']:
#         setCelltext(table,len(table.rows) - 1,1,data['ZJDH'],font_name_='Times New Roman')
#     setCelltext(table,len(table.rows) - 1,2,data['ZZDH'],font_name_='Times New Roman')
#     setCelltext(table, len(table.rows) - 1, 3, data['LZQLRMC'])
#     if re.search('村民小组', data['LZQLRMC']):
#         setCelltext(table, len(table.rows) - 1, 4, data['qz1'])
#     setCelltext(table, len(table.rows) - 1, 5, data['qz2'])
#     setCelltext(table, len(table.rows) - 1, 7, data['dkm'])


# 界址签章表
def generate_jzqz(data):
    doc = DocxTemplate(template_path / "界址签章表.docx")
    doc.render({"jzxdata": data})
    return doc


def generate_mjtj(zddmlist, zd_data):
    doc = Document(template_path / "村民小组地块面积统计表.docx")  # type: ignore
    doc.styles["Normal"].font.name = "方正仿宋_GBK"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "方正仿宋_GBK")
    table = doc.tables[0].cell(0, 0).tables[0]
    cell = doc.tables[0].cell(0, 0).tables[0].cell
    sum = 0
    for zddm in zddmlist:
        index = len(table.rows)
        table.add_row()
        table.rows[index].height = Cm(0.95)
        setCelltext(table, index, 0, index)
        setCelltext(table, index, 1, zddm)
        setCelltext(
            table, index, 2, str(zd_data[zd_data["ZDDM"] == zddm]["ZDMJ"].values[0])
        )
        sum += zd_data[zd_data["ZDDM"] == zddm]["ZDMJ"].values[0]

    table.add_row()
    table.rows[len(table.rows) - 1].height = Cm(1.49)
    setCelltext(table, len(table.rows) - 1, 0, "本村民小组总面积")
    cell(len(table.rows) - 1, 0).paragraphs[0].runs[0].font.size = Pt(14)
    cell(len(table.rows) - 1, 0).merge(cell(len(table.rows) - 1, 1))
    if type(sum) == "float":
        setCelltext(table, len(table.rows) - 1, 2, round(sum, 4))
    else:
        setCelltext(table, len(table.rows) - 1, 2, sum)
    return doc


def get_jxzx(jzd_jzsm, jzx_jzsm):
    if not jzx_jzsm.size:
        return ""
    result = f"{jzx_jzsm.QLRMC.values[0]}{jzx_jzsm.ZDDM.values[0]}\n"
    for _, row in jzx_jzsm.iterrows():
        if row["LZQLRMC"]:
            JZXLB = jzd_jzsm[jzd_jzsm["JZD_NEW"] == row["QSDH"]]["JZXLB"].values[0]
            H_JZXLB = jzd_jzsm[jzd_jzsm["JZD_NEW"] == row["ZZDH"]]["JZXLB"].values[0]
            QSDH = row["QSDH"]
            LZQLRMC = row["LZQLRMC"]
            ZJDH = f"过界址点{row['ZJDH']}" if row["ZJDH"] else ""
            ZZDH = row["ZZDH"]
            result += f"{QSDH}沿着{LZQLRMC}{JZXLB}{ZJDH}到{ZZDH}止，后为{H_JZXLB}。\n"
    return result


def generate_jzsm(zddmlist, qlr, jzd_data, jzx_data):
    doc = DocxTemplate(template_path / "界址说明表.docx")
    dwsm = ""  # 点位说明
    jxzx = ""  # 界线走向
    for zddm in zddmlist:
        log.info(f"{zddm}的界址信息")
        jzd_jzsm = jzd_data[jzd_data["ZDDM"] == zddm]
        jzx_jzsm = jzx_data[jzx_data["ZDDM"] == zddm]
        dwsmdict = {}
        if [v for v in jzd_jzsm.点位说明.values if v]:
            dwsm += f"{qlr}{zddm}\n"
        for _, row in jzd_jzsm.iterrows():
            if row["点位说明"]:
                if row["点位说明"] not in dwsmdict:
                    dwsmdict[row["点位说明"]] = row["JZD_NEW"]
                else:
                    dwsmdict[row["点位说明"]] = f"{dwsmdict[row['点位说明']]}、{row['JZD_NEW']}"
        for key, value in dwsmdict.items():
            dwsm += f"{value}位于{key}交界处。\n"

        jxzx += get_jxzx(jzd_jzsm, jzx_jzsm)
    doc.render({"DWSM": dwsm, "JXZX": jxzx})  
    return doc


# 宗地草图
def generate_zdct(jpg_path):
    doc = Document()
    doc.add_table(1, 1, style="Table Grid").rows[0].height = Cm(21.5)
    doc.tables[0].columns[0].width = Cm(15.21)
    run = doc.tables[0].cell(0, 0).paragraphs[0].add_run()
    run.add_picture(jpg_path, width=Cm(15.21), height=Cm(21.5))
    return doc


def generate_shb():
    doc = Document(template_path / "调查审核表.docx")  # type: ignore
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    par.paragraph_format.space_after = Pt(0)
    return doc


def generate_qjdc_(zd_data, jzd_data, jzx_data, savepath, jpg_zdct, control):
    
    for key, row in get_zd_data(zd_data, jzx_data, jzd_data).items():
        docxlist = []
        log.info(f"{row['QLRMC']}权籍调查表")
        jzjb_data = jzd_data[jzd_data.ZDDM.isin(zd_data[zd_data.QLRMC == key].ZDDM)]
        if control["head"]:
            head = generate_head(row["ZDDM"])
            docxlist.append(head)
        if control["dcb"]:
            dcb = generate_zd(row)
            docxlist.append(dcb)
        if control["jzjb"]:
            jzjb = generate_jzjb(jzjb_data)
            docxlist.append(jzjb)
        if control["qzb"]:
            qzb = generate_jzqz(row["jzx_data"])
            docxlist.append(qzb)
        if control["zdct"]:
            zdctdir = jpg_pathlist(jpg_zdct)
            if row["QLRMC"] not in zdctdir:
                log.err(f"{row['QLRMC']}缺失宗地草图")
                return f"{row['QLRMC']}缺失宗地草图"
            else:
                zdct = generate_zdct(zdctdir[row["QLRMC"]])
                docxlist.append(zdct)
        if control["mjtj"]:
            mjtj = generate_mjtj(row["ZDDM_list"], zd_data)
            docxlist.append(mjtj)
        if control["jzsm"]:
            jzsm = generate_jzsm(row["ZDDM_list"], key, jzd_data, jzx_data)
            docxlist.append(jzsm)
        if control["shb"]:
            shb = generate_shb()
            docxlist.append(shb)
        Djmod.compose_docx(docxlist, f"{savepath}\\{row['QLRMC']}权籍调查表.docx")
        yield f"{row['QLRMC']}权籍调查表"


if __name__ == "__main__":
    config = config.config
    gdb_path = config.gdb_path
