import logging, time, zipfile, os, traceback, sys
import pandas as pd
from multiprocessing import Pool
from functools import wraps
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docxcompose.composer import Composer
from pathlib import Path

def fileDF(directory_list: list[str]):
    df = pd.DataFrame(columns=["directory", "filename", "path", "type", "name"])
    for directory in directory_list:
        for root, _, files in os.walk(directory):
            if ".gdb" in root:
                continue
            for file in files:
                df.loc[df.shape[0]] = [
                    root,
                    file,
                    Path(root) / file,
                    file.split(".")[1],
                    file.split(".")[0],
                ]
    return df

def unzip(zip_path: str, unzip_path: str,filetype:str=''):
    '''解压文件'''
    with zipfile.ZipFile(zip_path, "r") as zip_file:
        namelist = []
        for info in zip_file.infolist():
            info.filename = info.filename.encode('cp437').decode('gbk')
            namelist.append(info.filename)
            zip_file.extract(info,unzip_path)
        if filetype == 'gdb':
            return zip_path
        return namelist

def zip_list(filelist: list[str|Path], zipname):
    # 多个文件压缩
    with zipfile.ZipFile(zipname, "w") as zip_file:
        for fpath in filelist:
            zip_file.write(fpath, arcname=str(fpath).split(os.sep)[-1])

def zipDir(dirpath, outFullName):
    """
    压缩指定文件夹
    :param dirpath: 目标文件夹路径
    :param outFullName: 压缩文件保存路径+xxxx.zip
    :return: 无
    """
    zip = zipfile.ZipFile(outFullName, "w", zipfile.ZIP_DEFLATED)
    for path, dirnames, filenames in os.walk(dirpath):
        # 去掉目标跟路径，只对目标文件夹下边的文件及文件夹进行压缩
        fpath = path.replace(dirpath, '')
 
        for filename in filenames:
            zip.write(os.path.join(path, filename), os.path.join(fpath, filename))
    zip.close()

def groupby(df: pd.DataFrame, by: list[str], agg: str):
    """agg:[
        'any','all','count','cov','first','idxmax',
        'idxmin','last','max','mean','median','min',
        'nunique','prod','quantile','sem','size',
        'skew','std','sum','var'
    ]
    """
    Aggfield = agg.upper()
    df2 = df.copy()
    df2[Aggfield] = ""
    by_df = pd.DataFrame(df2.groupby(by=by)[Aggfield].agg(agg))
    by_df.reset_index(inplace=True)
    return by_df

class Djlog:
    def __init__(self) -> None:
        # 日志输出
        logging.basicConfig(
            level=logging.INFO,
            filename=f"./log/{time.strftime('%Y%m%d', time.gmtime(time.time()))}.log",
            format="%(asctime)s %(filename)s:%(lineno)s %(message)s",
            datefmt="%Y-%m-%d %H:%M:%S",
            encoding="utf-8"
        )
        self.debug = logging.debug
        self.info = logging.info
        self.warning = logging.warning
        self.err = logging.error

def logErr(log: Djlog,):  # -> Callable[..., _Wrapped[Callable[..., Any], Any, Callable[..., Any], Any | str]]:# -> Callable[..., _Wrapped[Callable[..., Any], Any, Callable[..., Any], Any | str]]:# -> Callable[..., _Wrapped[Callable[..., Any], Any, Callable[..., Any], Any | str]]:# -> Callable[..., _Wrapped[Callable[..., Any], Any, Callable[..., Any], Any | str]]:# -> Callable[..., _Wrapped[Callable[..., Any], Any, Callable[..., Any], Any | str]]:# -> Callable[..., _Wrapped[Callable[..., Any], Any, Callable[..., Any], Any | str]]:# -> Callable[..., _Wrapped[Callable[..., Any], Any, Callable[..., Any], Any | str]]:
    # 错误日志输出
    def outwrapper(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            try:
                return func(*args, **kwargs)
            except BaseException as e:
                exc_type, exc_value, exc_traceback = sys.exc_info()
                error = "".join(
                    traceback.format_exception(exc_type, exc_value, exc_traceback)
                )
                log.err(error)
                return str("err")

        return wrapper

    return outwrapper

def compose_docx_file(files, output_file_path):
    """
    合并多个word文件到一个文件中
    :param files:待合并文件的列表
    :param output_file_path 新的文件路径
    :return:
    """
    composer = Composer(Document())
    n = 0
    for file in files:
        if not Path(file).exists():
            break
        doc = Document(file)
        if n < len(files) - 1:
            # 防止最后一个文档分页
            doc.add_page_break()
        composer.append(doc)
        n += 1

    composer.save(output_file_path)

def compose_docx(docxlist, output_file_path: str):
    """
    合并多个word文件到一个文件中
    :param files:待合并文件的列表
    :param output_file_path 新的文件路径
    :return:
    """
    composer = Composer(Document())
    n = 0
    for docx in docxlist:
        if n < len(docxlist) - 1:
            # 防止最后一个文档分页
            docx.add_page_break()
        composer.append(docx)
        n += 1
    composer.save(output_file_path)

def setCelltext(table_, row_, cell_, text_, fontname_="", font_size_=Pt(10.5)):
    """
      word单元格居中赋值
    Args:
        table_ (table): python-docx模块table对象
        row_ (number): 行号
        cell_ (number): 列号
        text_ (str/number): 需要赋值的内容
        fontname_ (str): 字体名称
        font_size_ ()
    """
    table_.rows[row_].cells[cell_].paragraphs[0].text = str(text_)
    table_.rows[row_].cells[
        cell_
    ].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table_.rows[row_].cells[cell_].paragraphs[
        0
    ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table_.rows[row_].cells[cell_].paragraphs[0].runs[0].font.size = font_size_
    if fontname_:
        table_.rows[row_].cells[cell_].paragraphs[0].runs[0].font.name = fontname_
        
def multiValueSplicing(data_path,fid,format_value,symbols='、'):
    """
    多值拼接
    :param data:数据
    :param fid:字段id
    :param format_value:格式化拼接值值
    :return:
    """
    df = pd.read_excel(data_path)
    data_dict = {}
    
    def dispose(row):
        if row[fid] not in data_dict.keys():
            data_dict[row[fid]] = row[format_value]
        else:
            data_dict[row[fid]] = f"{data_dict[row[fid]]}{symbols}{row[format_value]}"
    
    df.apply(dispose,axis=1)
    res_df = pd.DataFrame([{'id':k,'value':v} for k,v in data_dict.items() ])
    return res_df

class ProcessTask:
    def __init__(self,count:int):
        self.P = Pool(count)
        
    def exec(self,func,*args):
        self.P.apply_async(func,args=args)
        
    def loop(self,func,iterate,*args):
        for item in iterate:
            func(item,*args)
            
    def close(self):
        self.P.close()
        self.P.join()