import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from . import Djmod
from pathlib import Path
from . import config

template_path = Path(config.config.template_path)
log = Djmod.Djlog()
def jpg_pathlist(jpg_zdct):
    ct_path = {}
    for file in jpg_zdct.glob('*宗地草图.jpg'):
        ct_path[file.name[:-8]] = str(file)
    return ct_path

# docx文档赋值
def setCelltext(table_,
                row_,
                cell_,
                text_,
                paragraph_=0,
                run_=0,
                font_name_='',
                font_size_=Pt(10.5)):
    """_summary_
    word单元格居中赋值
  Args:
      table_ (table): python-docx模块table对象
      row_ (number): 行号
      cell_ (number): 列号
      text_ (str/number): 需要赋值的内容
      paragraph_ (int, optional): 段落号 Defaults to 0.
      run_ (int, optional): 样式号 Defaults to 0.
      font_name_ (str, optional): 字体名称. Defaults to ''.
      font_size_ (_type_, optional): 字体大小 Defaults to Pt(10.5).
  """
    table_.rows[row_].cells[cell_].paragraphs[paragraph_].add_run(str(text_))
    table_.rows[row_].cells[
        cell_].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table_.rows[row_].cells[cell_].paragraphs[
        paragraph_].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table_.rows[row_].cells[cell_].paragraphs[paragraph_].runs[
        run_].font.size = font_size_
    if font_name_:
        r = table_.rows[row_].cells[cell_].paragraphs[paragraph_].runs[run_]
        r.font.name = font_name_
        r._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_)

# 宗地数据格式化
def get_zd_data(zd_data,jzd_data,jzx_data):
    tidyup = {}
    n = 0
    # 宗地数据格式化
    for index, row in zd_data.iterrows():
        if row['QLRMC'] not in tidyup:
            n = 1
            tidyup[row['QLRMC']] = {
                **row,
                'NUM': n,
                'jzx_data': [],
                'jzd_data':{},
                'ZDDM_list': {row['ZDDM']}
            }
        else:
            n += 1
            tidyup[row['QLRMC']][
                'ZDDM'] = f"{tidyup[row['QLRMC']]['ZDDM']}、{row['ZDDM']}"
            tidyup[row['QLRMC']]['ZDMJ'] = tidyup[row['QLRMC']]['ZDMJ'] + row['ZDMJ']
            tidyup[row['QLRMC']][
                'BDCDYH'] = f"{tidyup[row['QLRMC']]['BDCDYH']}、{row['BDCDYH']}"
            if not row['TFH']:
                temp_tfh = set(row['TFH'].split('、'))
                tidyup_tfh = set(tidyup[row['QLRMC']]['TFH'].split('、'))
                over = tidyup_tfh | temp_tfh
                tidyup[row['QLRMC']]['TFH'] = '、'.join(over)
            tidyup[row['QLRMC']]['NUM'] = n
            tidyup[row['QLRMC']]['ZDDM_list'].add(row['ZDDM'])
        jzddf = jzd_data[jzd_data['ZDDM'] == row['ZDDM']]
        boundary_index = jzddf.INDEX.drop_duplicates()
        for _, jzd_row in jzddf.iterrows():
            try:
                
                if row['ZDDM'] not in tidyup[row['QLRMC']]['jzd_data']:
                    
                    tidyup[row['QLRMC']]['jzd_data'][row['ZDDM']] = [{**jzd_row}]
                else:
                    tidyup[row['QLRMC']]['jzd_data'][row['ZDDM']].append({**jzd_row})
            except KeyError:
                log.err(f"界址点与宗地不匹配: {row['ZDDM']}")
    for _, row in jzx_data.iterrows():
        try:
            xlzl = zd_data[zd_data['ZDDM'] == row['ZDDM']]
            tidyup[row['QLRMC']]['jzx_data'].append({
                'QLRMC':
                row['QLRMC'],
                'QSDH':
                row['QSDH'],
                'ZJDH':
                row['ZJDH'],
                'ZZDH':
                row['ZZDH'],
                'LZQLRMC':
                row['LZQLRMC'] if row['LZQLRMC'] else '',
                'qz1':
                f"详见{row['LZQLRMC']}界线认可书" if row['LZQLRMC'] else '',
                'qz2':
                f"详见{row['QLRMC']}界线认可书" if row['QLRMC'] else '',
                'dkm':
                xlzl.loc[xlzl.index[0], '飞地坐落'] if xlzl.loc[xlzl.index[0],
                                                            '飞地坐落'] else '',
            })
        except KeyError:
            log.err(f"界址线与宗地不匹配: {row['ZDDM']}")
    return tidyup

# 通用docx模板值替换
def template_docx(doc: Document, data_: dict, font_size=Pt(10.5)):
    for par in doc.paragraphs:
        for run in par.runs:
            Regex = re.search('@[A-Za-z]*@', par.text)
            if Regex:
                run.text = run.text.replace(
                    Regex.group(),
                    str(data_[Regex.group()[1:len(Regex.group()) - 1]]))

        # par.runs[0].font.name = fontname
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                Regex = re.search('@[A-Za-z]*@', cell.text)
                if Regex:
                    if data_[Regex.group()[1:len(Regex.group()) - 1]]:
                        cell.paragraphs[0].text = cell.paragraphs[
                            0].text.replace(
                                Regex.group(),
                                str(data_[Regex.group()[1:len(Regex.group()) -
                                                        1]]))
                        cell.paragraphs[0].runs[0].font.size = font_size
    return doc

def generate_head(zddm):
    doc = Document(template_path / '不动产权籍调查表.docx')
    doc.paragraphs[8].runs[2].text = doc.paragraphs[8].runs[2].text.replace('@ZDDM@', zddm)
    return doc

def generate_zd(data_):
    doc = Document(template_path / '宗地基本信息表.docx')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                Regex = re.search('@[A-Za-z]*@', cell.text)
                if Regex:
                    if data_[Regex.group()[1:len(Regex.group()) - 1]]:
                        cell.paragraphs[0].text = cell.paragraphs[
                            0].text.replace(
                                Regex.group(),
                                str(data_[Regex.group()[1:len(Regex.group()) -
                                                        1]]))
                        cell.paragraphs[0].space_after = Pt(0)
                    else:
                        cell.paragraphs[0].text = cell.paragraphs[
                            0].text.replace(Regex.group(), '')
                        for par in cell.paragraphs:
                            par.paragraph_format.space_after = Pt(0)
                    cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                    if Regex.group()[1:len(Regex.group()) - 1] in [
                            'ZJHM', 'LXDH', 'ZDDM', 'BDCDYH', 'TFH', 'ZDMJ',
                            'NUM'
                    ]:
                        cell.paragraphs[0].runs[
                            0].font.name = 'Times New Roman'
                    else:
                        cell.paragraphs[0].runs[0].font.name = u'方正仿宋_GBK'
                        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(
                            qn('w:eastAsia'), '方正仿宋_GBK')
                    if Regex.group()[1:len(Regex.group()) -
                                     1] in ['ZDMJ', 'NUM']:
                        for par in cell.paragraphs:
                            for run in par.runs:
                                run.font.size = Pt(7.5)
    return doc

def generate_jzjb(data_,jzd_data):
    doc = Document(template_path / '界址标示表.docx')
    doc.styles['Normal'].font.name = u'方正仿宋_GBK'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋_GBK')
    
    cell = doc.tables[0].cell
    
    countlist = [len(jzd_data.loc[jzd_data['ZDDM'] == zddm]) for zddm in data_]
    count = len(countlist) * 2
    for c in countlist:
        if c > 6:
            count += 12
        else:
            count += c * 2
    if count < 32:
        count = 32
    for i in range(count):
        doc.tables[0].add_row()
        doc.tables[0].rows[i + 3].height = Cm(0.5)

    for row_index in range(0, count, 2):
        if row_index <= count - 4:
            cell(row_index + 4, 0).merge(cell(row_index + 5, 0))
            cell(row_index + 4, 1).merge(cell(row_index + 5, 1))
        for colu_index in range(10):
            cell(row_index + 3,
                 colu_index + 2).merge(cell(row_index + 4, colu_index + 2))
    num = 0
    JZXLB_dict = {'沟渠':3,'道路':4,'田埂':5,'地埂':6,'山脊':7}
    JZXWZ_dict = {'内':8,'中':9,'外':10}
    for zddm,jdz_row in data_.items():
        if len(jdz_row) == 0:
            continue
        if len(jdz_row) <= 6:
            for row in jdz_row:
                setCelltext(doc.tables[0], num + 3, 0, row['JZD_NEW'])
                if row['JZXLB'] in JZXLB_dict:
                    setCelltext(doc.tables[0], num + 4, JZXLB_dict[row['JZXLB']], '√')
                else:
                    log.err(f"{zddm}-{row['JZD_NEW']}:界址线类别填写不正确")    

                if row['JZXWZ'] in JZXWZ_dict:
                    setCelltext(doc.tables[0], num + 4, JZXWZ_dict[row['JZXWZ']], '√')
                else:
                    log.err(f"{zddm}-{row['JZD_NEW']}:界址线位置填写不正确")  
                      
                if row['SM']:
                    setCelltext(doc.tables[0], num + 4, 11, row['SM'])
                num += 2

            setCelltext(doc.tables[0], num + 3, 0, 'J1')
            num += 2
        else:
            for i in range(6):
                if i < 3:
                    JZD_NEW =   jdz_row[i]['JZD_NEW']
                    JZXLB   =   jdz_row[i]['JZXLB']
                    JZXWZ   =   jdz_row[i]['JZXWZ']
                    SM      =   jdz_row[i]['SM']
                    setCelltext(doc.tables[0], num + 3, 0,JZD_NEW)
                    if JZXLB in JZXLB_dict:
                        setCelltext(doc.tables[0], num + 4, JZXLB_dict[JZXLB], '√')
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线类别填写不正确")
                        
                    if JZXWZ in JZXWZ_dict:
                        setCelltext(doc.tables[0], num + 4, JZXWZ_dict[JZXWZ], '√')
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线位置填写不正确")     

                    if SM :
                        setCelltext(doc.tables[0], num + 4, 11,SM )
                    num += 2
                elif i == 4:
                    setCelltext(doc.tables[0], num + 3, 0, '.....')
                    setCelltext(doc.tables[0], num + 3, 5, '√')
                    setCelltext(doc.tables[0], num + 3, 9, '√')
                    num += 2
                elif i > 4:
                    JZD_NEW =   jdz_row[-1]['JZD_NEW']
                    JZXLB   =   jdz_row[-1]['JZXLB']
                    JZXWZ   =   jdz_row[-1]['JZXWZ']
                    SM      =   jdz_row[-1]['SM']
                    setCelltext( doc.tables[0], num + 3, 0,JZD_NEW)
                    
                    if JZXLB in JZXLB_dict:
                        setCelltext(doc.tables[0], num + 4, JZXLB_dict[JZXLB], '√')
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线类别填写不正确")
                        
                    if JZXWZ in JZXWZ_dict:
                        setCelltext(doc.tables[0], num + 4, JZXWZ_dict[JZXWZ], '√')
                    else:
                        log.err(f"{zddm}-{JZD_NEW}:界址线位置填写不正确")     

                    if SM:
                        setCelltext(doc.tables[0], num + 4, 11, SM)
                    num += 2
            setCelltext(doc.tables[0], num + 3, 0, 'J1')
            num += 2
    par_ = doc.add_paragraph()
    par_.text = '说明：界址点标示示意情况详见集体土地所有权范围图。'
    par_.runs[0].font.size = Pt(10)
    return doc

# 界址签章表组件
def add_jzqz(table, data):
    table.add_row()
    setCelltext(table,len(table.rows) - 1,0, data['QSDH'],font_name_='Times New Roman')
    if data['ZJDH']:
        setCelltext(table,len(table.rows) - 1,1,data['ZJDH'],font_name_='Times New Roman')
    setCelltext(table,len(table.rows) - 1,2,data['ZZDH'],font_name_='Times New Roman')
    setCelltext(table, len(table.rows) - 1, 3, data['LZQLRMC'])
    if re.search('村民小组', data['LZQLRMC']):
        setCelltext(table, len(table.rows) - 1, 4, data['qz1'])
    setCelltext(table, len(table.rows) - 1, 5, data['qz2'])
    setCelltext(table, len(table.rows) - 1, 7, data['dkm'])

# 界址签章表
def generate_jzqz(data):
    doc = Document()
    doc.styles['Normal'].font.name = u'方正仿宋_GBK'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋_GBK')
    doc.add_table(3, 8, style='Table Grid')
    doc.tables[0].rows[0].height = Cm(1)
    doc.tables[0].rows[1].height = Cm(0.7)
    cell = doc.tables[0].cell
    setCelltext(doc.tables[0], 0, 0, '界址签章表')
    cell(0, 0).merge(cell(0, 7))

    setCelltext(doc.tables[0], 1, 0, '界址线')
    cell(1, 0).merge(cell(1, 2))

    setCelltext(doc.tables[0], 1, 3, '邻宗地')
    cell(1, 3).merge(cell(1, 4))

    setCelltext(doc.tables[0], 1, 5, '本宗地')
    setCelltext(doc.tables[0], 1, 6, '日期')
    setCelltext(doc.tables[0], 1, 7, '地块名')
    setCelltext(doc.tables[0], 2, 0, '起点号')
    setCelltext(doc.tables[0], 2, 1, '中间点')
    setCelltext(doc.tables[0], 2, 2, '终点号')
    setCelltext(doc.tables[0], 2, 3, '相邻宗地权利人（宗地号）')
    setCelltext(doc.tables[0], 2, 4, '指界人姓名（签章）')
    setCelltext(doc.tables[0], 2, 5, '指界人姓名（签章）')
    for row in data:
        add_jzqz(doc.tables[0], row)
    par_ = doc.add_paragraph()
    par_.text = '说明：协议书与上述界址点不一致的，以集体土地所有权范围图及界址点、界线的补充说明为准。'
    par_.runs[0].font.size = Pt(10)
    return doc

def generate_mjtj(zddmlist,zd_data):
    doc = Document(template_path / '村民小组地块面积统计表.docx')
    doc.styles['Normal'].font.name = u'方正仿宋_GBK'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋_GBK')
    table = doc.tables[0].cell(0, 0).tables[0]
    cell = doc.tables[0].cell(0, 0).tables[0].cell
    sum = 0
    for zddm in zddmlist:
        index = len(table.rows)
        table.add_row()
        table.rows[index].height = Cm(0.95)
        setCelltext(table, index, 0, index)
        setCelltext(table, index, 1, zddm)
        setCelltext(table, index, 2,
                    str(zd_data[zd_data['ZDDM'] == zddm]['ZDMJ'].values[0]))
        sum += zd_data[zd_data['ZDDM'] == zddm]['ZDMJ'].values[0]

    table.add_row()
    table.rows[len(table.rows) - 1].height = Cm(1.49)
    setCelltext(table, len(table.rows) - 1, 0, '本村民小组总面积')
    cell(len(table.rows) - 1, 0).paragraphs[0].runs[0].font.size = Pt(14)
    cell(len(table.rows) - 1, 0).merge(cell(len(table.rows) - 1, 1))
    if type(sum) == 'float':
        setCelltext(table, len(table.rows) - 1, 2, round(sum, 4))
    else:
        setCelltext(table, len(table.rows) - 1, 2, sum)
    return doc

def get_jxzx(jzd_jzsm,jzx_jzsm,row):
    if not jzx_jzsm.size:
        return ''
    result = f"{jzx_jzsm.QLRMC.values[0]}{jzx_jzsm.ZDDM.values[0]}\n"
    for index, row in jzx_jzsm.iterrows():
        if row['LZQLRMC']:
            JZXLB = jzd_jzsm[jzd_jzsm['JZD_NEW'] == row['QSDH']]['JZXLB'].values[0]
            QSDH = row['QSDH']
            LZQLRMC = row['LZQLRMC']
            ZJDH = f"过界址点{row['ZJDH']}" if row['ZJDH'] else ''
            ZZDH = row['ZZDH']
            result += f"{QSDH}沿着{LZQLRMC}{JZXLB}{ZJDH}到{ZZDH}止，后为曲线。\n"
    return result

def generate_jzsm(zddmlist, qlr,jzd_data,jzx_data):
    doc = Document(template_path / '界址说明表.docx')
    doc.styles['Normal'].font.name = u'方正仿宋_GBK'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋_GBK')
    dwsm = ''  # 点位说明
    jxzx = ''  # 界线走向
    for zddm in zddmlist:
        jzd_jzsm = jzd_data[jzd_data['ZDDM'] == zddm]
        jzx_jzsm = jzx_data[jzx_data['ZDDM'] == zddm]
        dwsm += f"{qlr}{zddm}\n"
        dwsmdict = {}
        for index, row in jzd_jzsm.iterrows():
            if row['点位说明']:
                if row['点位说明'] not in dwsmdict:
                    dwsmdict[row['点位说明']] = row['JZD_NEW']
                else:
                    dwsmdict[row['点位说明']] = f"{dwsmdict[row['点位说明']]}、{row['JZD_NEW']}"
        for key,value in dwsmdict.items():
            dwsm += f"{value}位于{key}交界处。\n"
        for index, row in jzx_jzsm.iterrows():
            jxzx += get_jxzx(jzd_jzsm,jzx_jzsm,row)
    template_docx(doc, {'DWSM': dwsm, 'JXZX': jxzx})
    return doc

# 宗地草图
def generate_zdct(jpg_path):
    doc = Document()
    doc.add_table(1, 1, style='Table Grid').rows[0].height = Cm(21.5)
    doc.tables[0].columns[0].width = Cm(15.21)
    run = doc.tables[0].cell(0, 0).paragraphs[0].add_run()
    run.add_picture(jpg_path, width=Cm(15.21), height=Cm(21.5))
    return doc

def generate_shb():
    doc = Document(template_path / '调查审核表.docx')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    par.paragraph_format.space_after = Pt(0)
    return doc   
   
def generate_qjdc(zd_data,jzd_data,jzx_data,savepath,jpg_zdct,control):
    for key, row in get_zd_data(zd_data,jzd_data,jzx_data).items():
        docxlist = []
        head = generate_head(row['ZDDM'])
        docxlist.append(head)
        if control['dcb']:
            dcb = generate_zd(row)
            docxlist.append(dcb)
        if control['jzjb']:
            jzjb = generate_jzjb(row['jzd_data'],jzd_data)
            docxlist.append(jzjb)
        if control['jzsm']:
            qzb = generate_jzqz(row['jzx_data'])
            docxlist.append(qzb)
        if control['zdct']:
            zdctdir = jpg_pathlist(jpg_zdct)
            if row['QLRMC'] not in jpg_pathlist():
                log.err(f"{row['QLRMC']}缺失宗地草图")
                return f"{row['QLRMC']}缺失宗地草图"
            else:
                zdct = generate_zdct(zdctdir[row['QLRMC']])
                docxlist.append(zdct)
        if control['mjtj']:
            mjtj = generate_mjtj(row['ZDDM_list'],zd_data)
            docxlist.append(mjtj)
        if control['jzsm']:
            jzsm = generate_jzsm(row['ZDDM_list'], key,jzd_data,jzx_data)
            docxlist.append(jzsm)
        if control['shb']:
            shb = generate_shb()
            docxlist.append(shb)
        Djmod.compose_docx(docxlist, f"{savepath}\\{row['QLRMC']}权籍调查表.docx")
        yield f"{row['QLRMC']}权属调查表！"
    
if __name__ == '__main__':
    
    config = config.config
    gdb_path = config.gdb_path

